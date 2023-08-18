from docx import Document
from docx.shared import RGBColor
from docx.shared import Inches
import os
import pandas as pd
import numpy as np
import warnings
warnings.filterwarnings('ignore')
import pickle
import fund_performance_calculator
import fund_performance_text_generator
import fund_performance_plot
import etf_data_preprocess
import etf_shares_and_value_calculator
import etf_text_generator

def set_heading_color(heading, color): # 此处是为了设置heading的颜色
    run = heading.runs[0]
    run.font.color.rgb = color

def add_bold_text(paragraph, text): # 此处是为了加粗字体
    run = paragraph.add_run(text)
    run.bold = True
    return run

# 此处作为global variable
input_date = '2023-08-14'

# 基金市场回顾
"""
以下是不含etf的基金部分
"""
# 四大类基金
qdii_fund_pickle = fund_performance_calculator.xml_to_pickles(r'../xml_file/qdii_xml.xml', "qdii") # 这里是qdii基金的xml文件
equity_fund_pickle = fund_performance_calculator.xml_to_pickles(r'../xml_file/equity_fund_xml.xml', "equity_fund") # 这里是偏股型基金的xml文件
index_fund_pickle = fund_performance_calculator.xml_to_pickles(r'../xml_file/index_fund_xml.xml', "index_fund") # 这里是指数型基金的xml文件
bond_fund_pickle = fund_performance_calculator.xml_to_pickles(r'../xml_file/bond_fund_xml.xml', "bond_fund") # 这里是偏债型基金的xml文件
#偏股基金
regular_pickle = fund_performance_calculator.xml_to_pickles(r'../xml_file/regular_stock_fund_xml.xml', "regular_stock_fund") # 这里普通股票型基金的xml文件
equity_skewed_pickle = fund_performance_calculator.xml_to_pickles(r'../xml_file/stocks_skewed_fund_xml.xml', "stocks_skewed_fund") # 这里是偏股混合型基金的xml文件
balanced_pickle = fund_performance_calculator.xml_to_pickles(r'../xml_file/balance_stock_fund_xml.xml', "balance_stock_fund") # 这里是平衡混合型基金的xml文件
# 偏债基金
convert_bond_pickle = fund_performance_calculator.xml_to_pickles(r'../xml_file/convert_bond_fund_xml.xml', "convert_bond_fund") # 这里是可转基的xml文件
secondary_bond_pickle = fund_performance_calculator.xml_to_pickles(r'../xml_file/secondary_bond_fund_xml.xml', "secondary_bond_fund") # 这里是二级债基的xml文件
bond_skewed_pickle = fund_performance_calculator.xml_to_pickles(r'../xml_file/bond_skewed_fund_xml.xml', "bond_skewed_fund") # 这里是偏债混合型基金的xml文件
regular_bond_pickle = fund_performance_calculator.xml_to_pickles(r'../xml_file/regular_bond_fund_xml.xml', "regular_bond_fund") # 这里是普通债券型基金的xml文件
interest_rate_pickle = fund_performance_calculator.xml_to_pickles(r'../xml_file/interest_rate_fund_xml.xml', "interest_rate_fund") # 这里是利率债基的xml文件
pure_bond_pickle = fund_performance_calculator.xml_to_pickles(r'../xml_file/pure_bond_fund_xml.xml', "pure_bond_fund") # 这里是纯债基的xml文件

# 四类基金结果，分别是qdii、偏股型基金、指数型基金、偏债型基金
four_fund_output = fund_performance_calculator.calculate_fund_performance(qdii_fund_pickle, equity_fund_pickle,\
                                                                          index_fund_pickle, bond_fund_pickle, input_date)
qdii_last_week_returns = four_fund_output[0]
qdii_new_year_returns = four_fund_output[1]
qdii_positive_percentage = four_fund_output[2]
equity_last_week_returns = four_fund_output[3]
equity_new_year_returns = four_fund_output[4]
equity_positive_percentage = four_fund_output[5]
index_last_week_returns = four_fund_output[6]
index_new_year_returns = four_fund_output[7]
index_positive_percentage = four_fund_output[8]
bond_last_week_returns = four_fund_output[9]
bond_new_year_returns = four_fund_output[10]
bond_positive_percentage = four_fund_output[11]
# 全场基金回顾文本生成
four_fund_text = fund_performance_text_generator.create_fund_review_document(qdii_last_week_returns, qdii_new_year_returns, qdii_positive_percentage,\
                                                                             equity_last_week_returns, equity_new_year_returns, equity_positive_percentage,\
                                                                             index_last_week_returns, index_new_year_returns, index_positive_percentage,\
                                                                             bond_last_week_returns, bond_new_year_returns, bond_positive_percentage)
# 全场基金回顾的两个图
# 第一个是周收益和年收益图
fund_performance_plot.four_category_fund_plot(qdii_last_week_returns, qdii_new_year_returns,\
                                              equity_last_week_returns, equity_new_year_returns,\
                                              index_last_week_returns, index_new_year_returns,\
                                              bond_last_week_returns, bond_new_year_returns)
# 第二个是正收益占比图
fund_performance_plot.four_category_fund_positive_percentage_plot(qdii_positive_percentage, equity_positive_percentage,\
                                                                  index_positive_percentage, bond_positive_percentage)
import png_merge
# 这里是要把两个图片合并在一起
png_merge.merge_images_horizontally(r'../png_file/四类基金收益.png', r'../png_file/四类基金正收益.png')
# 传入数据
stock_fund_output = fund_performance_calculator.calculate_equity_fund_performance(regular_pickle, equity_skewed_pickle, balanced_pickle, input_date)
stock_last_week_returns = stock_fund_output[0]
hybrid_last_week_returns = stock_fund_output[3]
balanced_last_week_returns = stock_fund_output[6]
# 普通股票，偏股混合，平衡混合
# 偏股基金回顾文本
stock_fund_text = fund_performance_text_generator.create_equity_fund_review_document(equity_last_week_returns, equity_new_year_returns,\
                                                                                     stock_last_week_returns, hybrid_last_week_returns,\
                                                                                     balanced_last_week_returns)
# 偏股基金的图，含周收益和年收益
fund_performance_plot.stock_skewed_fund_plot(stock_fund_output[0], stock_fund_output[1],\
                                             stock_fund_output[3], stock_fund_output[4],\
                                             stock_fund_output[6], stock_fund_output[7])
# 偏债基金回顾文本
# 传入参数
bond_fund_output = fund_performance_calculator.calculate_bond_fund_performance(convert_bond_pickle, secondary_bond_pickle,\
                                                                               bond_skewed_pickle, regular_bond_pickle,\
                                                                               interest_rate_pickle, pure_bond_pickle, input_date)
bond_average_last_week_returns = four_fund_output[9]
bond_average_year_returns = four_fund_output[10]
convertible_last_week_returns = bond_fund_output[0]
secondary_last_week_returns = bond_fund_output[3]
hybrid_bond_last_week_returns = bond_fund_output[6]
bond_last_week_returns = bond_fund_output[9]
interest_rate_last_week_returns = bond_fund_output[12]
pure_bond_last_week_returns = bond_fund_output[15]
bond_fund_text = fund_performance_text_generator.create_bond_fund_review_document(bond_average_last_week_returns, bond_average_year_returns,\
                                                                                  convertible_last_week_returns, secondary_last_week_returns,
                                                                                  hybrid_bond_last_week_returns, bond_last_week_returns,
                                                                                  interest_rate_last_week_returns, pure_bond_last_week_returns)
# 偏债基金画图
fund_performance_plot.bond_skewed_fund_plot(bond_fund_output[0], bond_fund_output[1],\
                                            bond_fund_output[3], bond_fund_output[4],\
                                            bond_fund_output[6], bond_fund_output[7],\
                                            bond_fund_output[9], bond_fund_output[10],\
                                            bond_fund_output[12], bond_fund_output[13],\
                                            bond_fund_output[15], bond_fund_output[16])

"""
以下是基金的etf部分
"""
sectors_chinese = ['军工', '光伏', '稀土', '低碳环保', '机械设备', '新材料', '周期', '价值', '农业养殖', '一带一路',\
                   '新能源汽车', '红利低波', '科创创业50', '医药', '中证500', '成长', '创业板指', '金融地产', '证券',\
                   '银行', '沪深300', 'MSCI', '上证50', '恒生中国企业', '恒生系列指数', '恒生科技', '通信电子',\
                   '恒生指数', '科技', '沪港深', '科创50', '消费', '食品饮料', '游戏传媒', '人工智能', '大数据', '计算机', '芯片半导体']

sectors = ['JUNGONG', 'GUANGFU', 'XITU', 'DITANHUANBAO', 'JIXIESHEBEI', 'XINCAILIAO', 'ZHOUQI', 'JIAZHI', 'NONGYEYANGZHI', 'YIDAIYILU',\
           'XINNENGYUANQICHE', 'HONGLIDIBO', 'KECHUANGCHUANGYE50', 'YIYAO', 'ZHONGZHENG500', 'CHENGZHANG', 'CHUANGYEBANZHI', 'JINRONGDICHAN', 'ZHENGQUAN',\
           'YINHANG', 'HUSHEN300', 'MSCI', 'SHANGZHENG50', 'HENGSHENGZHONGGUOQIYE', 'HENGSHENGXILIEZHISHU', 'HENGSHENGKEJI', 'TONGXINDIANZI',\
           'HENGSHENGZHISHU', 'KEJI', 'HUGANGSHEN', 'KECHUANG50', 'XIAOFEI', 'SHIPINYINLIAO', 'YOUXICHUANMEI', 'RENGONGZHINENG', 'DASHUJU', 'JISUANJI',\
           'XINPIANBANDAOTI']

sector_dict = dict(zip(sectors, sectors_chinese))

# iterate所有板块
for key, value in sector_dict.items():
    # 根据板块名称构造对应的数据处理对象
    # data_obj = etf_data_preprocess.XmlData(sector, r'../xml_file/etf_DASHUJU.xml')
    etf_sector_path = r'../xml_file'
    etf_file_name = f'etf_{key}.xml'
    data_obj = etf_data_preprocess.XmlData(value, f'{etf_sector_path}/{etf_file_name}')
    # 调用数据处理对象的方法进行数据处理
    data_obj.to_final_value_pickle()
    data_obj.to_final_portion_compare_pickle()

# for sector in sectors:
#     # 根据板块名称构造对应的数据处理对象
#     data_obj = etf_data_preprocess.XmlData(sector, f"etf_{sector}.xml")
#     # 调用数据处理对象的方法进行数据处理
#     data_obj.to_final_value_pickle()
#     data_obj.to_final_portion_compare_pickle()
# 其实是替代了下列的代码，生成了后续所需要pickle文件
# df_jungong = etf_data_preprocess.XmlData('军工', "etf_jungong.xml")
# df_jungong.to_final_value_pickle()
# df_jungong.to_final_portion_compare_pickle()

if __name__ == "__main__":
    # 定义35个不同的 sectors
    sectors = ['军工', '光伏', '稀土', '低碳环保', '机械设备', '新材料', '周期', '价值', '农业养殖', '一带一路',\
               '新能源汽车', '红利低波', '科创创业50', '医药', '中证500', '成长', '创业板指', '金融地产', '证券',\
               '银行', '沪深300', 'MSCI', '上证50', '恒生中国企业', '恒生系列指数', '恒生科技', '通信电子',\
               '恒生指数', '科技', '沪港深', '科创50', '消费', '食品饮料', '游戏传媒', '人工智能', '大数据', '计算机',\
               '芯片半导体']

    # sectors = ['JUNGONG', 'GUANGFU', 'XITU', 'DITANHUANBAO', 'JIXIESHEBEI', 'XINCAILIAO', 'ZHOUQI', 'JIAZHI','NONGYEYANGZHI', 'YIDAIYILU', \
    #            'XINNENGYUANQICHE', 'HONGLIDIBO', 'KECHUANGCHUANGYE50', 'YIYAO', 'ZHONGZHENG500', 'CHENGZHANG', 'CHUANGYEBANZHI', 'JINRONGDICHAN', 'ZHENGQUAN', \
    #            'KUANJI', 'YINHANG', 'HUSHEN300', 'MSCI', 'SHANGZHENG50', 'HENGSHENGZHONGGUOQIYE', 'HENGSHENGXILIEZHISHU', 'HENGSHENGKEJI', 'TONGXINDIANZI', \
    #            'HENGSHENGZHISHU', 'KEJI', 'HUGANGSHEN', 'KECHUANG50', 'XIAOFEI', 'SHIPINYINLIAO', 'YOUXICHUANMEI', 'RENGONGZHINENG', 'DASHUJU', 'JISUANJI', \
    #            'XINPIANBANDAOTI']

    # sectors_dict = {'JUNGONG':'军工', 'GUANGFU':'光伏', 'XITU':'稀土', 'DITANHUANBAO':'低碳环保','JIXIESHEBEI':'机械设备', 'XINCAILIAO':'新材料', 'ZHOUQI':'周期',\
    #                 'JIAZHI':'价值', 'NONGYEYANGZHI':'农业养殖',}

    # 定义每个sector对应的文件路径信息
    data_info = {}
    for sector in sectors:
        # 此处是测试所用，可以取消comment
        # file_path_value = f"{sector}_value.pickle"
        # file_path_portion = f"{sector}_portion.pickle"
        value_file_name = f'{sector}_value.pickle'
        portion_file_name = f'{sector}_portion.pickle'
        etf_path = r'../pickle_file'
        file_path_value = f"{etf_path}/{value_file_name}"
        file_path_portion = f"{etf_path}/{portion_file_name}"
        data_info[sector] = {"file_path_value": file_path_value, "file_path_portion": file_path_portion}

    date = input_date
    modes = ["shares_change", "value_change", "top_etf", "calculate_shares_sum"]

    all_results = {}  # 存储所有 sector 的计算结果

    for sector, info in data_info.items():
        sector_results = {}  # 存储当前 sector 的计算结果

        for mode in modes:
            file_path = info["file_path_value"] if mode == "value_change" else info["file_path_portion"]

            if os.path.exists(file_path):
                calculator = etf_shares_and_value_calculator.ProcessedEtfPickleCalculator(sector, file_path, mode, date)

                if mode == "shares_change":
                    result = calculator.calculator_portion_change()
                elif mode == "value_change":
                    result = calculator.calculator_value_change()
                elif mode == "top_etf":
                    result = calculator.find_top_three_funds()
                elif mode == "calculate_shares_sum":
                    result = calculator.calculate_shares_sum()
                else:
                    raise ValueError("Invalid mode. 请重新输入.")

                print(f"Results for {sector} - Mode: {mode}")
                print(result)

                # 将当前mode的计算结果存储在当前sector的结果字典中
                sector_results[mode] = result

            else:
                print(f"{sector} 板块的文件没找到 - Mode: {mode}")

        # 将当前sector的结果字典存储在大的结果字典中
        all_results[sector] = sector_results

    # 打印所有 sector 的结果
    print()
    print("All Results:")
    print('*' * 15000)
    print(all_results)
    print('*' * 15000)
    print()
    # all_results.to_pickle('主题ETF计算结果.pickle')
    # 这里是保存成为中间结果，可以去看
    with open(r'../pickle_file/主题ETF计算结果.pickle', 'wb') as files:
        pickle.dump(all_results, files)

# 经过测试，import没办法直接运行python文件，一下代码是为了运行两个py文件所写，且这两行代码只能放在此处
exec(open('etf_find_top_five_2.py', encoding='utf-8').read())
exec(open('etf_fill_excel.py', encoding='utf-8').read())

# 此处是为了生成etf部分的文本
etf_text = etf_text_generator.create_etf_review_text(r'../pickle_file/etf_for_excel.pickle')

"""
接下来是股票市场
"""
import stock_market_data_preprocess
import stock_market_calculator
import stock_market_plot
import stock_market_text_generator

stock_file = r'../xml_file/stock_market_ten_index.xml' # 此处是海内外股市指数的xml文件
shenwan_file = r'../xml_file/shenwan_industry_index.xml' # 此处是申万31个行业的xml文件

# 此处stock_df和shenwan_df分别都是处理好的数据
stock_df = stock_market_data_preprocess.stock_market_xml_preprocessing(stock_file)
shenwan_df = stock_market_data_preprocess.shenwan_index_xml_preprocessing(shenwan_file)

# 以下带output的variables都是计算好的数据，text相应的文本，title是被用作标题的文本
stock_text_output = stock_market_calculator.stock_index_text_calculator(stock_df, input_date)
stock_domestic_text = stock_market_text_generator.domestic_stock_market_text(stock_text_output)
stock_oversea_text = stock_market_text_generator.oversea_stock_market_text(stock_text_output)
stock_domestic_title = stock_market_text_generator.domestic_stock_market_title_text(stock_text_output)

shenwan_output = stock_market_calculator.shenwan_index_weekly_return_calculator(shenwan_df, input_date)
shenwan_text = stock_market_text_generator.shenwan_text(shenwan_output)
shenwan_title = stock_market_text_generator.shenwan_title_text(shenwan_output)

# 此处是画股市行业的图，海内外股市以及申万行业均在其中
stock_plot_output = stock_market_calculator.stock_index_weekly_return_calculator(stock_df, input_date)
stock_market_plot.stock_market_ten_index_plot(stock_plot_output)
stock_market_plot.shenwan_index_plot(shenwan_output)

"""
接下来是宽基指数
"""
import broad_based_index_allinone
broad_index_file = r'../xml_file/broad_based_index_xml.xml' # 此处是宽基指数的xml文件
broad_index_output = broad_based_index_allinone.broad_based_index_preprocess(broad_index_file)
pe_now = broad_based_index_allinone.find_now(broad_index_output[0], input_date)
pb_now = broad_based_index_allinone.find_now(broad_index_output[1], input_date)
pe_ptl = broad_based_index_allinone.pepb_ratio_percentile(broad_index_output[0], input_date)
pb_ptl = broad_based_index_allinone.pepb_ratio_percentile(broad_index_output[1], input_date)
pe_range = broad_based_index_allinone.pe_range(pe_ptl)
broad_index_text = broad_based_index_allinone.broad_index_text_generation(pe_now, pe_ptl, input_date, pe_range)
broad_based_index_allinone.broad_index_plot(pe_range, pe_now, pe_ptl, pb_now, pb_ptl)

import new_issued_fund
new_issued_fund_df = new_issued_fund.new_fund_data_cleaning(r'../xml_file/latest_issued_fund.xml') # 此处是本周新发基金xml
previous_issued_fund_df = new_issued_fund.new_fund_data_cleaning(r'../xml_file/last_week_issued_fund.xml') # 此处是上周新发基金xml

# 这里要把上周的新发基金，放在第二个传入参数
new_issued_title = new_issued_fund.new_issued_fund_title(new_issued_fund_df, previous_issued_fund_df)
new_issued_text = new_issued_fund.new_issued_fund_text(new_issued_fund_df, previous_issued_fund_df)
new_issued_fund.new_issued_fund_graph(new_issued_fund_df)

# 全场基金回顾文本
doc = Document()
doc.styles['Normal'].font.name = '宋体'
title_text = '「基金市场周报」' + stock_domestic_title
heading_title_text = doc.add_heading(title_text, level = 1)
set_heading_color(heading_title_text, RGBColor(0, 0, 0))
main_conclusion_heading = doc.add_heading('主要结论',level=1)
set_heading_color(main_conclusion_heading, RGBColor(255, 0, 0))
main_conclusion_text = stock_domestic_title + shenwan_title
heading_main_conclusion = doc.add_heading(main_conclusion_text, level=1)
set_heading_color(heading_main_conclusion, RGBColor(0, 0, 0))
doc.add_paragraph(stock_domestic_text + shenwan_text)
doc.add_paragraph()
doc.add_paragraph() # 学sample周报里面的空两行
heading_new_issued_title = doc.add_heading(new_issued_title,level=1)
set_heading_color(heading_new_issued_title, RGBColor(0, 0, 0))
doc.add_paragraph(new_issued_text)
doc.add_paragraph()
doc.add_paragraph() # 学sample周报里面的空两行

# 全场基金回顾文本
heading_fund = doc.add_heading('基金市场回顾',level=1)
set_heading_color(heading_fund, RGBColor(255, 0, 0))
paragraph_all_fund = doc.add_paragraph()
add_bold_text(paragraph_all_fund, '全市场基金回顾:')
paragraph_all_fund.add_run(four_fund_text)
doc.add_picture(r'../png_file/合并.png', width=Inches(6))
# 偏股基金回顾文本
paragraph_stock_fund = doc.add_paragraph()
add_bold_text(paragraph_stock_fund, "偏股基金回顾：")
paragraph_stock_fund.add_run(stock_fund_text)
doc.add_picture(r'../png_file/偏股基金收益.png', width=Inches(6))
# 偏债基金回顾文本
doc.add_paragraph(bond_fund_text)
paragraph_bond_fund = doc.add_paragraph()
add_bold_text(paragraph_bond_fund, "偏债基金回顾：")
paragraph_bond_fund.add_run(bond_fund_text)
doc.add_picture(r'../png_file/偏债基金收益.png', width=Inches(6))
paragraph_etf = doc.add_paragraph()

add_bold_text(paragraph_etf, "权益ETF表现：") # 添加带粗体的文本
paragraph_etf.add_run(etf_text)


# 股市市场回顾文本
heading_stock = doc.add_heading('股票市场回顾',level=1)
set_heading_color(heading_stock, RGBColor(255, 0, 0))
paragraph_domestic = doc.add_paragraph() # 添加一个段落
add_bold_text(paragraph_domestic, "国内股市：") # 添加带粗体的文本
doc.add_paragraph(stock_domestic_text + shenwan_text)
paragraph_oversea = doc.add_paragraph() # 添加一个段落
add_bold_text(paragraph_oversea, "海外股市：") # 添加带粗体的文本
paragraph_oversea.add_run(stock_oversea_text)
doc.add_picture(r'../png_file/海内外股市指数.png', width=Inches(6))
doc.add_picture(r'../png_file/申万行业指数.png', width=Inches(6))

# 宽基指数部分回顾文本
heading_broad_index = doc.add_heading('市场估值情况',level=1)
set_heading_color(heading_broad_index, RGBColor(255, 0, 0))
paragraph_broad_index = doc.add_paragraph()
add_bold_text(paragraph_broad_index, "宽基指数PE：") # 添加带粗体的文本
paragraph_broad_index.add_run(broad_index_text)
doc.add_picture(r'../png_file/宽基指数.png', width=Inches(6))

# 新发基金文本回顾
heading_new = doc.add_heading('新发基金',level=1)
set_heading_color(heading_new, RGBColor(255, 0, 0))
paragraph_new = doc.add_paragraph()
red_text_run = paragraph_new.add_run('本周新发基金盘点')
red_text_run.font.color.rgb = RGBColor(255, 0, 0)  # 设置为红色
doc.add_paragraph(new_issued_text)
doc.add_picture(r'../png_file/新发基金.png', width=Inches(6))


# 保存成为word文件
doc.save(r'../word_file/first_word.docx')

# 用来提示文件生成成功
print()
print()
print('！！！！！！！！！！！！文档合成结束！！！！！！！！！！')



